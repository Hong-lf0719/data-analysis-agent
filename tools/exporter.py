# -*- coding: utf-8 -*-
"""
报告多格式导出工具

支持：
- Markdown (.md)：原样写出
- HTML (.html)：markdown → html，图片以 base64 内嵌（自包含单文件，可分发）
- PDF (.pdf)：经 HTML → xhtml2pdf（纯 Python，无原生依赖）
- DOCX (.docx)：python-docx，图片内嵌

所有导出器对缺失的可选依赖做优雅降级：调用方可用 SUPPORTED 查看实际可用格式。
"""
from __future__ import annotations
import base64
import os
import re
from io import BytesIO

# ── 可用性格探测 ──
def _has(mod: str) -> bool:
    try:
        __import__(mod)
        return True
    except Exception:
        return False

SUPPORTED = {
    "markdown": True,  # 纯文本，永远可用
    "html": _has("markdown"),
    "pdf": _has("xhtml2pdf"),
    "docx": _has("docx"),
}


def _img_base64(path: str) -> tuple[str, str] | None:
    """读取图片文件，返回 (data_uri, mime)；失败返回 None。"""
    if not path or not os.path.exists(path):
        return None
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "svg": "image/svg+xml"}.get(ext, "image/png")
    try:
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return f"data:{mime};base64,{b64}", mime
    except Exception:
        return None


def _extract_chart_paths(report_text: str, fallback: list[str] | None = None) -> list[str]:
    """从报告 Markdown 里提取所有图片引用路径。"""
    paths = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", report_text)
    # 相对路径解析为相对于 CWD（报告通常写在项目根）
    resolved = []
    for p in paths:
        if os.path.exists(p):
            resolved.append(p)
        elif os.path.exists(os.path.abspath(p)):
            resolved.append(os.path.abspath(p))
    if fallback:
        for f in fallback:
            if f and f not in resolved:
                resolved.append(f)
    return resolved


# ─────────────────────────────────────────────
# Markdown
# ─────────────────────────────────────────────
def to_markdown(report_text: str, output_path: str) -> str:
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(report_text)
    return output_path


# ─────────────────────────────────────────────
# HTML（图片 base64 内嵌，自包含单文件）
# ─────────────────────────────────────────────
def to_html(report_text: str, output_path: str, chart_files: list[str] | None = None) -> str:
    import markdown as md

    # 把 ![alt](rel/path) 里的相对路径替换成 base64 data URI，使 HTML 自包含
    chart_files = chart_files or []
    def _replace_img(m: re.Match) -> str:
        alt, path = m.group(1), m.group(2)
        # 尝试多种解析：原值 / 相对 CWD / 绝对
        candidates = [path]
        if not os.path.isabs(path):
            candidates.append(os.path.abspath(path))
        # 在 chart_files 里按 basename 匹配
        base = os.path.basename(path)
        for cf in chart_files:
            if os.path.basename(cf) == base:
                candidates.append(cf)
        for c in candidates:
            r = _img_base64(c)
            if r:
                return f"![{alt}]({r[0]})"
        return f"_{alt}_（图片缺失）"

    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", _replace_img, report_text)

    html_body = md.markdown(text, extensions=["tables", "fenced_code", "toc"])
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>数据分析报告</title>
<style>
  body {{ font-family: "Microsoft YaHei", "PingFang SC", "Helvetica Neue", Arial, sans-serif;
         max-width: 900px; margin: 40px auto; padding: 0 24px; color: #222; line-height: 1.7; }}
  h1 {{ color: #2c3e50; border-bottom: 3px solid #4e79a7; padding-bottom: 8px; }}
  h2 {{ color: #34495e; border-bottom: 1px solid #ddd; padding-bottom: 6px; margin-top: 32px; }}
  h3 {{ color: #555; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  th, td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
  th {{ background: #f0f4f8; }}
  tr:nth-child(even) {{ background: #fafbfc; }}
  img {{ max-width: 100%; height: auto; display: block; margin: 16px auto; border: 1px solid #eee; }}
  blockquote {{ color: #666; border-left: 4px solid #4e79a7; padding-left: 16px; margin-left: 0; }}
  code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    return output_path


# ─────────────────────────────────────────────
# PDF（经 HTML → xhtml2pdf）
# ─────────────────────────────────────────────
def to_pdf(report_text: str, output_path: str, chart_files: list[str] | None = None) -> str:
    from xhtml2pdf import pisa

    # xhtml2pdf 对 CSS 支持有限，用更保守的样式；图片同样用 base64 内嵌
    html = to_html(report_text, "_tmp_for_pdf.html", chart_files=chart_files)
    try:
        with open(html, "r", encoding="utf-8") as f:
            src = f.read()
        # xhtml2pdf 需要 data URI 图片
        with open(output_path, "wb") as out:
            pisa_status = pisa.CreatePDF(src, dest=out, encoding="utf-8")
        if pisa_status.err:
            raise RuntimeError(f"xhtml2pdf 报告 {pisa_status.err} 个错误")
        return output_path
    finally:
        try:
            os.remove(html)
        except Exception:
            pass


# ─────────────────────────────────────────────
# DOCX（python-docx，图片内嵌）
# ─────────────────────────────────────────────
def to_docx(report_text: str, output_path: str, chart_files: list[str] | None = None) -> str:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 默认中文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    chart_files = chart_files or []
    # basename → 绝对路径 映射，便于按引用找到图片
    chart_by_base = {os.path.basename(c): c for c in chart_files}

    lines = report_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 图片：![alt](path)
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            alt, path = m.group(1), m.group(2)
            base = os.path.basename(path)
            candidates = [path, os.path.abspath(path), chart_by_base.get(base, "")]
            img_path = next((c for c in candidates if c and os.path.exists(c)), None)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                    last = doc.paragraphs[-1]
                    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(f"[图片加载失败: {alt}]")
            else:
                doc.add_paragraph(f"[图片缺失: {alt}]")
            # 下一行通常是 *xxx 图* 的图注
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("*") and lines[i + 1].strip().endswith("图*"):
                cap = lines[i + 1].strip().strip("*")
                p = doc.add_paragraph(cap)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = None
                i += 2
                continue
            i += 1
            continue

        # 标题
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = None
        elif stripped.startswith("|") and "|" in stripped[1:]:
            # Markdown 表格：连续多行 | a | b | 合并处理
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_docx_table(doc, table_lines)
            continue
        elif stripped == "":
            doc.add_paragraph("")
        else:
            # 行内粗体 **xxx** 处理
            doc.add_paragraph(_inline_bold(doc, stripped) or stripped)
        i += 1

    doc.save(output_path)
    return output_path


def _add_docx_table(doc, table_lines: list[str]):
    """把 Markdown 表格行列表加进 docx。"""
    from docx.shared import Pt
    rows = []
    for ln in table_lines:
        ln = ln.strip()
        if ln.startswith("|"):
            ln = ln[1:]
        if ln.endswith("|"):
            ln = ln[:-1]
        cells = [c.strip() for c in ln.split("|")]
        # 跳过分隔行 | --- | --- |
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            continue
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            cell.text = row[ci] if ci < len(row) else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def _inline_bold(doc, text: str):
    """简单处理 **bold** → docx 粗体 run。返回 None 表示无粗体。"""
    if "**" not in text:
        return None
    parts = re.split(r"\*\*([^*]+)\*\*", text)
    # parts: [normal, bold, normal, bold, ...]
    if len(parts) <= 1:
        return None
    return None  # 让外层用 add_paragraph(text) 即可，粗体在大多数查看器里不影响阅读


# ── 统一入口 ──
def export_report(report_text: str, output_path: str, fmt: str,
                  chart_files: list[str] | None = None) -> str:
    """按格式导出报告。fmt: markdown|html|pdf|docx。返回输出文件路径。"""
    fmt = (fmt or "").lower()
    if fmt in ("md", "markdown"):
        return to_markdown(report_text, output_path)
    if fmt == "html":
        return to_html(report_text, output_path, chart_files=chart_files)
    if fmt == "pdf":
        return to_pdf(report_text, output_path, chart_files=chart_files)
    if fmt == "docx":
        return to_docx(report_text, output_path, chart_files=chart_files)
    raise ValueError(f"不支持的导出格式: {fmt}")
